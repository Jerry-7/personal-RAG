# Personal RAG - DOCX 解析器
"""
Word 文档解析器模块

使用 python-docx 提取 .docx 文件的文本内容。
保留段落和表格结构信息。
"""

from app.services.parser.base import BaseParser, ParsedDocument


class DocxParser(BaseParser):
    """
    Word (.docx) 文件解析器。

    提取段落文本和表格内容，保持文档结构。
    不支持旧版 .doc 格式（需要 LibreOffice 转换）。
    """

    @property
    def supported_extensions(self) -> list[str]:
        """支持 .docx 文件。"""
        return ["docx"]

    def parse(self, file_path: str) -> ParsedDocument:
        """
        解析 .docx 文件，提取文本内容。

        Args:
            file_path: .docx 文件的本地路径

        Returns:
            ParsedDocument: 包含提取的文本
                - text: 完整文本（段落间用换行分隔）
                - metadata: {title, author, paragraph_count}

        Raises:
            FileNotFoundError: 文件不存在
            ValueError: 文件不是有效的 .docx 格式
        """
        # 延迟导入，避免未安装时的导入错误
        from docx import Document as DocxDocument

        doc = DocxDocument(file_path)

        # 提取文档属性
        props = doc.core_properties
        metadata = {
            "title": props.title or "",
            "author": props.author or "",
            "paragraph_count": len(doc.paragraphs),
        }

        text_parts: list[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # 也提取表格中的文本
        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_texts.append(cell.text.strip())
                if row_texts:
                    text_parts.append(" | ".join(row_texts))

        return ParsedDocument(
            text="\n\n".join(text_parts),
            metadata=metadata,
            file_type="docx",
        )
